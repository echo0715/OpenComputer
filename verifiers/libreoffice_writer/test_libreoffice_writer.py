"""
Test LibreOffice Writer verifier endpoints in a live E2B sandbox.

Covers:
  - Help/usage output
  - Error cases (LibreOffice not running, bad args)
  - ODF file parsing endpoints (ODT file, no UNO needed)
  - UNO live endpoints with DOCX file (text, paragraphs, formatting, tables, etc.)
  - Composite check-* endpoints — positive and negative cases
  - Heading detection, word count, search, images

Usage:
    python verifiers/libreoffice_writer/test_libreoffice_writer.py
"""

import json
import sys
import time
import traceback
from pathlib import Path

from dotenv import load_dotenv
from e2b_desktop import Sandbox
from e2b.sandbox.commands.command_handle import CommandExitException

load_dotenv()

VERIFIER_LOCAL = Path(__file__).parent / "libreoffice_writer.py"
VERIFIER_REMOTE = "/home/user/verifiers/libreoffice_writer.py"
V = f"python3 {VERIFIER_REMOTE}"

TEST_ODT_PATH = "/home/user/test_verifier.odt"
TEST_DOCX_PATH = "/home/user/test_verifier.docx"

# ── ODT fixture (stdlib ZIP+XML) ────────────────────────────────────────

CREATE_ODT_SCRIPT = r'''
import zipfile

CONTENT_XML = """<?xml version="1.0" encoding="UTF-8"?>
<office:document-content
  xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0"
  xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0"
  xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0"
  xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0"
  xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0"
  office:version="1.2">
  <office:automatic-styles>
    <style:style style:name="bold" style:family="text">
      <style:text-properties fo:font-weight="bold"/>
    </style:style>
  </office:automatic-styles>
  <office:body>
    <office:text>
      <text:h text:style-name="Heading_20_1" text:outline-level="1">Introduction</text:h>
      <text:p text:style-name="Standard">This is the first paragraph of the document.</text:p>
      <text:p text:style-name="Standard">This is the second paragraph with some <text:span text:style-name="bold">bold text</text:span> in it.</text:p>
      <text:h text:style-name="Heading_20_2" text:outline-level="2">Details</text:h>
      <text:p text:style-name="Standard">A detailed paragraph about the topic.</text:p>
      <table:table table:name="TestTable">
        <table:table-column table:number-columns-repeated="2"/>
        <table:table-row>
          <table:table-cell><text:p>Name</text:p></table:table-cell>
          <table:table-cell><text:p>Score</text:p></table:table-cell>
        </table:table-row>
        <table:table-row>
          <table:table-cell><text:p>Alice</text:p></table:table-cell>
          <table:table-cell><text:p>90</text:p></table:table-cell>
        </table:table-row>
        <table:table-row>
          <table:table-cell><text:p>Bob</text:p></table:table-cell>
          <table:table-cell><text:p>75</text:p></table:table-cell>
        </table:table-row>
      </table:table>
      <text:p text:style-name="Standard">Paragraph between tables.</text:p>
      <table:table table:name="InventoryTable">
        <table:table-column table:number-columns-repeated="4"/>
        <table:table-row>
          <table:table-cell><text:p>Item</text:p></table:table-cell>
          <table:table-cell><text:p>Quantity</text:p></table:table-cell>
          <table:table-cell><text:p>Price</text:p></table:table-cell>
          <table:table-cell><text:p>Total</text:p></table:table-cell>
        </table:table-row>
        <table:table-row>
          <table:table-cell><text:p>Widget</text:p></table:table-cell>
          <table:table-cell><text:p>10</text:p></table:table-cell>
          <table:table-cell><text:p>25.50</text:p></table:table-cell>
          <table:table-cell><text:p>255.00</text:p></table:table-cell>
        </table:table-row>
        <table:table-row>
          <table:table-cell><text:p>Gadget</text:p></table:table-cell>
          <table:table-cell><text:p>5</text:p></table:table-cell>
          <table:table-cell><text:p>50.00</text:p></table:table-cell>
          <table:table-cell><text:p></text:p></table:table-cell>
        </table:table-row>
        <table:table-row>
          <table:table-cell><text:p></text:p></table:table-cell>
          <table:table-cell table:number-columns-repeated="2"><text:p></text:p></table:table-cell>
          <table:table-cell><text:p>505.00</text:p></table:table-cell>
        </table:table-row>
      </table:table>
    </office:text>
  </office:body>
</office:document-content>
"""

MANIFEST_XML = """<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
  <manifest:file-entry manifest:media-type="application/vnd.oasis.opendocument.text" manifest:full-path="/"/>
  <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="content.xml"/>
</manifest:manifest>
"""

path = "/home/user/test_verifier.odt"
with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
    z.writestr("mimetype", "application/vnd.oasis.opendocument.text")
    z.writestr("content.xml", CONTENT_XML)
    z.writestr("META-INF/manifest.xml", MANIFEST_XML)

print("OK")
'''

# ── DOCX fixture (python-docx) ──────────────────────────────────────────
# Creates a document with headings, paragraphs, a table, and formatting

CREATE_DOCX_SCRIPT = r'''
import subprocess, sys
subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"],
                      stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Heading 1
h1 = doc.add_heading("Project Report", level=1)

# Normal paragraph
p1 = doc.add_paragraph("This document contains a project report with multiple sections.")

# Bold paragraph
p2 = doc.add_paragraph()
run = p2.add_run("This entire paragraph is bold.")
run.bold = True

# Heading 2
doc.add_heading("Data Summary", level=2)

# Table
table = doc.add_table(rows=3, cols=3)
table.style = "Table Grid"
headers = ["Item", "Quantity", "Price"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [["Widget", "10", "25.00"], ["Gadget", "5", "50.00"]]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

# Another heading
doc.add_heading("Conclusion", level=2)

# Conclusion paragraph
doc.add_paragraph("The project was completed successfully with all deliverables met on time.")

doc.save("/home/user/test_verifier.docx")
print("OK")
'''

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

passed = 0
failed = 0
errors: list[str] = []


class CmdResult:
    def __init__(self, exit_code: int, stdout: str, stderr: str):
        self.exit_code = exit_code
        self.stdout = stdout
        self.stderr = stderr


def run(sandbox: Sandbox, cmd: str, timeout: int = 30) -> dict | list:
    r = run_raw(sandbox, cmd, timeout)
    if r.exit_code != 0 and not r.stdout.strip():
        return {"error": f"exit_code={r.exit_code} stderr={r.stderr[:300]}"}
    try:
        return json.loads(r.stdout)
    except json.JSONDecodeError:
        return {"error": f"Invalid JSON: {r.stdout[:300]}"}


def run_raw(sandbox: Sandbox, cmd: str, timeout: int = 30) -> CmdResult:
    try:
        result = sandbox.commands.run(f"{V} {cmd}", timeout=timeout)
        return CmdResult(result.exit_code, result.stdout, result.stderr)
    except CommandExitException as e:
        return CmdResult(e.exit_code, e.stdout, e.stderr)


def run_shell(sandbox: Sandbox, cmd: str, timeout: int = 60) -> CmdResult:
    try:
        result = sandbox.commands.run(cmd, timeout=timeout)
        return CmdResult(result.exit_code, result.stdout, result.stderr)
    except CommandExitException as e:
        return CmdResult(e.exit_code, e.stdout, e.stderr)


def check(name: str, condition: bool, detail: str = ""):
    global passed, failed
    if condition:
        passed += 1
        print(f"  PASS  {name}")
    else:
        failed += 1
        msg = f"  FAIL  {name}"
        if detail:
            msg += f"  -- {detail}"
        print(msg)
        errors.append(f"{name}: {detail}")


def is_valid_json(stdout: str) -> bool:
    try:
        json.loads(stdout)
        return True
    except (json.JSONDecodeError, TypeError):
        return False


# ---------------------------------------------------------------------------
# Test groups
# ---------------------------------------------------------------------------

def test_help(sandbox: Sandbox):
    print("\n=== Help ===")
    result = run_raw(sandbox, "--help")
    check("help exits 0", result.exit_code == 0, f"got exit_code={result.exit_code}")
    check("help mentions commands", "Commands:" in result.stdout, result.stdout[:100])


def test_errors_lo_not_running(sandbox: Sandbox):
    print("\n=== Errors (LibreOffice not running) ===")
    for cmd, name in [("text", "text"), ("paragraphs", "paragraphs"), ("tables", "tables")]:
        data = run(sandbox, cmd)
        check(f"{name} returns error when LO down", "error" in data, str(data)[:100])


def test_errors_bad_args(sandbox: Sandbox):
    print("\n=== Errors (bad args) ===")
    result = run_raw(sandbox, "check-paragraph-text")
    check("missing arg exits 1", result.exit_code == 1)
    check("missing arg valid JSON", is_valid_json(result.stdout), result.stdout[:100])

    result = run_raw(sandbox, "nonexistent-command")
    check("unknown cmd exits 1", result.exit_code == 1)
    check("unknown cmd valid JSON", is_valid_json(result.stdout), result.stdout[:100])


def test_odt_file_parsing(sandbox: Sandbox):
    """ODF parsing endpoints (no UNO needed)."""
    print("\n=== ODT File Parsing ===")

    data = run(sandbox, f"parse-text {TEST_ODT_PATH}")
    check("parse-text returns text", "text" in data and "error" not in data, str(data)[:150])
    if "error" not in data:
        check("parse-text has Introduction", "Introduction" in data.get("text", ""), str(data.get("text", ""))[:100])

    data = run(sandbox, f"parse-paragraphs {TEST_ODT_PATH}")
    check("parse-paragraphs returns list", len(data.get("paragraphs", [])) > 0, str(data)[:150])
    if "error" not in data:
        paras = data["paragraphs"]
        # Check heading detection
        headings = [p for p in paras if p.get("heading")]
        check("parse-paragraphs finds headings", len(headings) >= 2, f"headings={len(headings)}")
        check("parse-paragraphs first heading = Introduction",
              headings[0]["text"] == "Introduction" if headings else False,
              str(headings[0]) if headings else "no headings")

    data = run(sandbox, f"parse-tables {TEST_ODT_PATH}")
    check("parse-tables returns tables", len(data.get("tables", [])) > 0, str(data)[:150])
    if "error" not in data:
        tbls = data["tables"]
        check("parse-tables has 2 tables", len(tbls) == 2,
              f"got {len(tbls)} tables: {[t.get('name') for t in tbls]}")

        # TestTable: 3 rows x 2 cols (header + Alice + Bob)
        t0 = tbls[0] if tbls else {}
        check("parse-tables first table is TestTable", t0.get("name") == "TestTable",
              str(t0.get("name")))
        check("TestTable has 3 rows", t0.get("rows") == 3, f"rows={t0.get('rows')}")
        check("TestTable has 2 cols", t0.get("cols") == 2, f"cols={t0.get('cols')}")
        if t0.get("data") and len(t0["data"]) >= 3:
            check("TestTable header row", t0["data"][0] == ["Name", "Score"],
                  str(t0["data"][0]))
            check("TestTable Alice row", t0["data"][1] == ["Alice", "90"],
                  str(t0["data"][1]))
            check("TestTable Bob row", t0["data"][2] == ["Bob", "75"],
                  str(t0["data"][2]))

        # InventoryTable: 4 rows x 4 cols (header + Widget + Gadget + total)
        t1 = tbls[1] if len(tbls) > 1 else {}
        check("parse-tables second table is InventoryTable", t1.get("name") == "InventoryTable",
              str(t1.get("name")))
        check("InventoryTable has 4 rows", t1.get("rows") == 4, f"rows={t1.get('rows')}")
        check("InventoryTable has 4 cols", t1.get("cols") == 4, f"cols={t1.get('cols')}")
        if t1.get("data") and len(t1["data"]) >= 4:
            check("InventoryTable header", t1["data"][0] == ["Item", "Quantity", "Price", "Total"],
                  str(t1["data"][0]))
            check("InventoryTable Widget row", t1["data"][1][0] == "Widget",
                  str(t1["data"][1]))
            check("InventoryTable Widget price", t1["data"][1][2] == "25.50",
                  str(t1["data"][1][2]))
            # Empty cell in Gadget total
            check("InventoryTable Gadget total empty", t1["data"][2][3] == "",
                  repr(t1["data"][2][3]))
            # Row with number-columns-repeated: empty cells should be expanded
            check("InventoryTable total row has 4 cols",
                  len(t1["data"][3]) == 4, f"cols={len(t1['data'][3])}")
            check("InventoryTable grand total", t1["data"][3][3] == "505.00",
                  str(t1["data"][3][3]))

    data = run(sandbox, "parse-text /nonexistent/file.odt")
    check("parse missing file returns error", "error" in data, str(data)[:100])


def test_uno_basic(sandbox: Sandbox):
    """Basic UNO endpoints with DOCX loaded."""
    print("\n=== UNO Basic (DOCX) ===")

    data = run(sandbox, "text")
    check("text returns text", "text" in data and "error" not in data, str(data)[:150])
    if "error" not in data:
        check("text has content", len(data.get("text", "")) > 0, f"length={data.get('length')}")

    data = run(sandbox, "doc-info")
    check("doc-info has title", "title" in data and "error" not in data, str(data)[:150])

    data = run(sandbox, "page-count")
    check("page-count returns dict", isinstance(data, dict) and "error" not in data, str(data)[:100])
    if "error" not in data:
        check("page-count >= 1", data.get("page_count", 0) >= 1, f"page_count={data.get('page_count')}")


def test_uno_paragraphs(sandbox: Sandbox):
    """Paragraph reading via UNO."""
    print("\n=== UNO Paragraphs ===")

    data = run(sandbox, "paragraphs")
    check("paragraphs returns list", "paragraphs" in data and "error" not in data, str(data)[:150])
    if "error" not in data:
        paras = data["paragraphs"]
        check("paragraphs count >= 5", len(paras) >= 5, f"count={len(paras)}")
        # First paragraph should be "Project Report" heading
        check("first para = Project Report", "Project Report" in paras[0].get("text", ""),
              f"text={paras[0].get('text')}")

    data = run(sandbox, "paragraph-format 0")
    check("paragraph-format returns dict", isinstance(data, dict) and "error" not in data, str(data)[:200])
    if "error" not in data:
        check("paragraph-format has style", "style" in data, str(data.keys()))
        check("paragraph-format has alignment", "alignment" in data, str(data.keys()))
        check("paragraph-format has bold", "bold" in data, str(data.keys()))


def test_uno_tables(sandbox: Sandbox):
    """Table endpoints via UNO."""
    print("\n=== UNO Tables ===")

    data = run(sandbox, "tables")
    check("tables returns dict", isinstance(data, dict) and "error" not in data, str(data)[:150])
    if "error" not in data:
        check("tables count >= 1", data.get("count", 0) >= 1, f"count={data.get('count')}")
        if data.get("count", 0) >= 1:
            t0 = data["tables"][0]
            check("table has name", "name" in t0, str(t0))
            check("table has rows", t0.get("rows", 0) >= 1, f"rows={t0.get('rows')}")
            check("table has cols", t0.get("cols", 0) >= 1, f"cols={t0.get('cols')}")

    # Get data from first table by name
    if "error" not in data and data.get("count", 0) > 0:
        table_name = data["tables"][0]["name"]
        tdata = run(sandbox, f"table-data {table_name}")
        check("table-data has data", "data" in tdata and "error" not in tdata, str(tdata)[:200])
        if "error" not in tdata and "data" in tdata:
            check("table first cell = Item", tdata["data"][0][0] == "Item",
                  f"first_cell={tdata['data'][0][0]}")
            # Check numeric value handling
            if len(tdata["data"]) >= 2 and len(tdata["data"][1]) >= 2:
                check("table data row has Widget", tdata["data"][1][0] == "Widget",
                      f"cell={tdata['data'][1][0]}")
                # Quantity cell should be numeric (10)
                check("table quantity is numeric",
                      tdata["data"][1][1] == 10 or tdata["data"][1][1] == "10",
                      f"val={tdata['data'][1][1]} type={type(tdata['data'][1][1])}")

    # table-data by index (no name)
    tdata = run(sandbox, "table-data")
    check("table-data default index 0", "data" in tdata or "error" in tdata, str(tdata)[:100])

    # table-data for nonexistent table
    tdata = run(sandbox, "table-data NonExistentTableZZZ")
    check("table-data nonexistent returns error", "error" in tdata, str(tdata)[:100])


def test_uno_misc(sandbox: Sandbox):
    """Misc UNO endpoints: search, images, page-style, bookmarks, headers."""
    print("\n=== UNO Misc ===")

    data = run(sandbox, 'search "project"')
    check("search returns dict", isinstance(data, dict) and "error" not in data, str(data)[:150])
    if "error" not in data:
        check("search found=bool", isinstance(data.get("found"), bool), str(data))

    data = run(sandbox, "images")
    check("images returns dict", isinstance(data, dict) and "error" not in data, str(data)[:100])
    if "error" not in data:
        check("images has count", "count" in data, str(data.keys()))

    data = run(sandbox, "page-style")
    check("page-style returns dict", isinstance(data, dict) and "error" not in data, str(data)[:200])
    if "error" not in data:
        check("page-style has width", "width" in data, str(data.keys()))
        check("page-style has orientation", "orientation" in data, str(data.keys()))

    data = run(sandbox, "bookmarks")
    check("bookmarks returns dict", isinstance(data, dict) and "error" not in data, str(data)[:100])

    data = run(sandbox, "headers-footers")
    check("headers-footers returns dict", isinstance(data, dict) and "error" not in data, str(data)[:200])


def test_checks_positive(sandbox: Sandbox):
    """Composite check-* endpoints — positive cases."""
    print("\n=== Checks (positive) ===")

    # check-text-contains
    data = run(sandbox, 'check-text-contains "project report"')
    check("check-text-contains found", data.get("contains") is True, str(data)[:150])
    check("check-text-contains has snippet", data.get("snippet") is not None, str(data.get("snippet")))

    # check-paragraph-text (first para = "Project Report")
    paras = run(sandbox, "paragraphs")
    if "error" not in paras:
        first_text = paras["paragraphs"][0]["text"]
        data = run(sandbox, f'check-paragraph-text 0 "{first_text}"')
        check("check-paragraph-text match", data.get("match") is True, str(data)[:150])

    # check-heading-exists
    data = run(sandbox, 'check-heading-exists "Project Report"')
    check("check-heading-exists found", data.get("exists") is True, str(data)[:150])

    data = run(sandbox, 'check-heading-exists "Data Summary" 2')
    check("check-heading-exists level=2 found", data.get("exists") is True, str(data)[:150])

    # check-table-exists (UNO)
    tables = run(sandbox, "tables")
    if "error" not in tables and tables.get("count", 0) > 0:
        tname = tables["tables"][0]["name"]
        data = run(sandbox, f"check-table-exists {tname}")
        check("check-table-exists found", data.get("exists") is True, str(data)[:100])

        # check-table-cell (UNO) — header
        data = run(sandbox, f"check-table-cell {tname} 0 0 Item")
        check("check-table-cell header match", data.get("match") is True, str(data)[:150])

        # check-table-cell (UNO) — data cell
        data = run(sandbox, f"check-table-cell {tname} 1 0 Widget")
        check("check-table-cell Widget match", data.get("match") is True, str(data)[:150])

        # check-table-cell (UNO) — numeric cell
        data = run(sandbox, f"check-table-cell {tname} 1 1 10")
        check("check-table-cell numeric match", data.get("match") is True, str(data)[:150])

        # check-table-cell (UNO) — negative: wrong value
        data = run(sandbox, f"check-table-cell {tname} 0 0 WrongValue")
        check("check-table-cell wrong value false", data.get("match") is False, str(data)[:150])

        # check-table-cell (UNO) — negative: out of range row
        data = run(sandbox, f"check-table-cell {tname} 999 0 Item")
        check("check-table-cell out of range row", data.get("match") is False, str(data)[:150])

    # ── File-based table checks (no UNO needed) ──
    # Save the current document as ODT first for file-based checks
    saved_odt = "/home/user/test_saved.odt"
    save_result = run_shell(sandbox, f"""python3 -c "
import uno
ctx = uno.getComponentContext()
resolver = ctx.ServiceManager.createInstanceWithContext('com.sun.star.bridge.UnoUrlResolver', ctx)
rctx = resolver.resolve('uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext')
smgr = rctx.ServiceManager
desktop = smgr.createInstanceWithContext('com.sun.star.frame.Desktop', rctx)
doc = desktop.getCurrentComponent()
from com.sun.star.beans import PropertyValue
props = (PropertyValue('FilterName', 0, 'writer8', 0),)
doc.storeToURL('file://{saved_odt}', props)
print('SAVED')
" """, timeout=30)
    file_saved = "SAVED" in save_result.stdout

    if file_saved:
        # check-table-exists-file — positive
        data = run(sandbox, f"check-table-exists-file Table1 {saved_odt}")
        # Note: DOCX table names may differ in ODT export. Accept whatever name it has.
        if data.get("exists") is not True and "available" in data:
            # Use whatever table name was exported
            available = data.get("available", [])
            if available:
                actual_name = available[0]
                data = run(sandbox, f"check-table-exists-file {actual_name} {saved_odt}")
                check("check-table-exists-file found (exported name)", data.get("exists") is True, str(data)[:100])
        else:
            check("check-table-exists-file found", data.get("exists") is True, str(data)[:100])

        # check-table-exists-file — negative
        data = run(sandbox, f"check-table-exists-file NonExistentTableZZZ {saved_odt}")
        check("check-table-exists-file not found", data.get("exists") is False, str(data)[:100])

        # check-table-cell-file — get actual table name first
        pt = run(sandbox, f"parse-tables {saved_odt}")
        if "error" not in pt and pt.get("count", 0) > 0:
            ft_name = pt["tables"][0]["name"]
            ft_data = pt["tables"][0].get("data", [])

            # Positive: check header cell
            if ft_data and len(ft_data[0]) > 0:
                expected_val = ft_data[0][0]
                data = run(sandbox, f"check-table-cell-file {ft_name} 0 0 {expected_val} {saved_odt}")
                check("check-table-cell-file header match", data.get("match") is True, str(data)[:150])

            # Positive: check data cell
            if len(ft_data) > 1 and len(ft_data[1]) > 0:
                expected_val = ft_data[1][0]
                data = run(sandbox, f"check-table-cell-file {ft_name} 1 0 {expected_val} {saved_odt}")
                check("check-table-cell-file data match", data.get("match") is True, str(data)[:150])

            # Negative: wrong value
            data = run(sandbox, f"check-table-cell-file {ft_name} 0 0 TotallyWrongValueZZZ {saved_odt}")
            check("check-table-cell-file wrong value", data.get("match") is False, str(data)[:150])

            # Negative: out of range
            data = run(sandbox, f"check-table-cell-file {ft_name} 999 0 test {saved_odt}")
            check("check-table-cell-file out of range", data.get("match") is False, str(data)[:150])

    # Also test file-based checks on the ODT fixture (TestTable and InventoryTable)
    data = run(sandbox, f"check-table-exists-file TestTable {TEST_ODT_PATH}")
    check("check-table-exists-file TestTable", data.get("exists") is True, str(data)[:100])

    data = run(sandbox, f"check-table-exists-file InventoryTable {TEST_ODT_PATH}")
    check("check-table-exists-file InventoryTable", data.get("exists") is True, str(data)[:100])

    data = run(sandbox, f"check-table-cell-file TestTable 0 0 Name {TEST_ODT_PATH}")
    check("check-table-cell-file TestTable Name", data.get("match") is True, str(data)[:150])

    data = run(sandbox, f"check-table-cell-file TestTable 1 0 Alice {TEST_ODT_PATH}")
    check("check-table-cell-file TestTable Alice", data.get("match") is True, str(data)[:150])

    data = run(sandbox, f"check-table-cell-file TestTable 1 1 90 {TEST_ODT_PATH}")
    check("check-table-cell-file TestTable 90", data.get("match") is True, str(data)[:150])

    data = run(sandbox, f"check-table-cell-file TestTable 2 0 Bob {TEST_ODT_PATH}")
    check("check-table-cell-file TestTable Bob", data.get("match") is True, str(data)[:150])

    data = run(sandbox, f"check-table-cell-file InventoryTable 0 0 Item {TEST_ODT_PATH}")
    check("check-table-cell-file InventoryTable header", data.get("match") is True, str(data)[:150])

    data = run(sandbox, f"check-table-cell-file InventoryTable 1 2 25.50 {TEST_ODT_PATH}")
    check("check-table-cell-file InventoryTable price", data.get("match") is True, str(data)[:150])

    # Negative: wrong table name
    data = run(sandbox, f"check-table-cell-file WrongTable 0 0 test {TEST_ODT_PATH}")
    check("check-table-cell-file wrong table", data.get("match") is False, str(data)[:150])

    # Nonexistent file
    data = run(sandbox, "check-table-exists-file TestTable /nonexistent.odt")
    check("check-table-exists-file missing file", "error" in data, str(data)[:100])

    # check-word-count
    data = run(sandbox, "check-word-count 5")
    check("check-word-count in_range", data.get("in_range") is True, str(data)[:100])

    # check-file-exists
    data = run(sandbox, f"check-file-exists {TEST_DOCX_PATH}")
    check("check-file-exists true", data.get("exists") is True, str(data)[:100])

    # check-file-saved
    data = run(sandbox, "check-file-saved")
    check("check-file-saved has key", "saved" in data or "error" in data, str(data)[:100])

    # check-image-count (no images in our test doc)
    data = run(sandbox, "check-image-count 0")
    check("check-image-count 0 match", data.get("match") is True, str(data)[:100])


def test_checks_negative(sandbox: Sandbox):
    """Composite check-* endpoints — negative cases."""
    print("\n=== Checks (negative) ===")

    data = run(sandbox, 'check-text-contains "xyzzy_nonexistent_text_12345"')
    check("check-text-contains false", data.get("contains") is False, str(data)[:100])
    check("check-text-contains snippet=None", data.get("snippet") is None, str(data))

    data = run(sandbox, 'check-paragraph-text 0 "WrongText12345"')
    check("check-paragraph-text mismatch", data.get("match") is False, str(data)[:100])

    data = run(sandbox, 'check-heading-exists "NonexistentHeading12345"')
    check("check-heading-exists false", data.get("exists") is False, str(data)[:100])

    data = run(sandbox, "check-table-exists NonExistentTable12345")
    check("check-table-exists false", data.get("exists") is False, str(data)[:100])

    # check-table-cell negative — wrong value via UNO
    tables = run(sandbox, "tables")
    if "error" not in tables and tables.get("count", 0) > 0:
        tname = tables["tables"][0]["name"]
        data = run(sandbox, f"check-table-cell {tname} 0 0 NonexistentValue12345")
        check("check-table-cell wrong value false", data.get("match") is False, str(data)[:150])

    data = run(sandbox, "check-file-exists /nonexistent/file.odt")
    check("check-file-exists false", data.get("exists") is False, str(data)[:100])

    data = run(sandbox, "check-image-count 99")
    check("check-image-count mismatch", data.get("match") is False, str(data)[:100])

    data = run(sandbox, "check-word-count 999999")
    check("check-word-count out of range", data.get("in_range") is False, str(data)[:100])


def test_all_commands_return_json(sandbox: Sandbox):
    """Every CLI command should output valid JSON."""
    print("\n=== JSON validity (all commands) ===")

    no_arg_cmds = ["text", "paragraphs", "doc-info", "page-count", "tables",
                   "images", "page-style", "bookmarks", "headers-footers"]
    for cmd in no_arg_cmds:
        result = run_raw(sandbox, cmd)
        valid = is_valid_json(result.stdout)
        check(f"{cmd} valid JSON", valid,
              f"exit={result.exit_code} stdout={result.stdout[:80]}" if not valid else "")

    arg_cmds = [
        ("paragraph-format", "0"),
        ("search", '"project"'),
        ("parse-text", TEST_ODT_PATH),
        ("parse-paragraphs", TEST_ODT_PATH),
        ("parse-tables", TEST_ODT_PATH),
        ("check-text-contains", '"test"'),
        ("check-paragraph-count", "5"),
        ("check-paragraph-text", '0 "test"'),
        ("check-paragraph-style", '0 "Heading"'),
        ("check-paragraph-formatted", "0 true"),
        ("check-heading-exists", '"test"'),
        ("check-word-count", "1"),
        ("check-table-exists-file", f"TestTable {TEST_ODT_PATH}"),
        ("check-table-cell-file", f"TestTable 0 0 Name {TEST_ODT_PATH}"),
        ("check-file-exists", TEST_DOCX_PATH),
        ("check-file-saved", ""),
        ("check-image-count", "0"),
    ]

    for cmd, arg in arg_cmds:
        full_cmd = f"{cmd} {arg}".strip()
        result = run_raw(sandbox, full_cmd)
        valid = is_valid_json(result.stdout)
        check(f"{full_cmd} valid JSON", valid,
              f"exit={result.exit_code} stdout={result.stdout[:80]}" if not valid else "")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def launch_libreoffice(sandbox: Sandbox, file_path: str) -> bool:
    """Launch LibreOffice Writer with UNO socket, opening the given file."""
    try:
        sandbox.commands.run("pkill -9 -f soffice", timeout=5)
    except (CommandExitException, Exception):
        pass
    try:
        sandbox.commands.run("pkill -9 -f oosplash", timeout=5)
    except (CommandExitException, Exception):
        pass
    time.sleep(2)

    launch_script = f'''#!/bin/bash
soffice --writer \
  '--accept=socket,host=localhost,port=2002;urp;' \
  --norestore --nologo \
  {file_path} \
  > /tmp/lo.log 2>&1
'''
    sandbox.files.write("/tmp/launch_lo.sh", launch_script)
    sandbox.commands.run("chmod +x /tmp/launch_lo.sh", timeout=3)
    sandbox.commands.run("/tmp/launch_lo.sh", background=True)

    check_script = '''
import uno, time
for attempt in range(5):
    try:
        ctx = uno.getComponentContext()
        resolver = ctx.ServiceManager.createInstanceWithContext(
            "com.sun.star.bridge.UnoUrlResolver", ctx
        )
        resolved = resolver.resolve(
            "uno:socket,host=localhost,port=2002;urp;StarOffice.ComponentContext"
        )
        smgr = resolved.ServiceManager
        desktop = smgr.createInstanceWithContext("com.sun.star.frame.Desktop", resolved)
        doc = desktop.getCurrentComponent()
        if doc is not None:
            print("READY")
            break
        else:
            print("NODOC")
            break
    except Exception as e:
        if attempt < 4:
            time.sleep(1)
        else:
            print(f"ERR:{e}")
'''
    sandbox.files.write("/tmp/check_uno.py", check_script)

    for i in range(30):
        try:
            r = sandbox.commands.run("python3 /tmp/check_uno.py", timeout=10)
            if "READY" in r.stdout:
                print(f"  LibreOffice UNO ready (attempt {i+1})")
                return True
            if "NODOC" in r.stdout:
                print(f"  UNO connected but no doc yet (attempt {i+1})")
        except (CommandExitException, Exception) as e:
            if i > 15:
                print(f"  attempt {i+1}: {str(e)[:80]}")
        time.sleep(2)

    try:
        r = sandbox.commands.run("tail -20 /tmp/lo.log", timeout=3)
        print(f"  [lo log] {r.stdout[-300:]}")
    except (CommandExitException, Exception):
        pass
    return False


def main():
    global passed, failed

    print("=" * 60)
    print("LibreOffice Writer Verifier Test Suite")
    print("=" * 60)

    print("\nCreating sandbox from desktop-all-apps...")
    sandbox = Sandbox.create(template="desktop-all-apps", timeout=600)

    try:
        # Upload verifier
        print(f"Uploading {VERIFIER_LOCAL} -> {VERIFIER_REMOTE}")
        sandbox.commands.run("mkdir -p /home/user/verifiers")
        with open(VERIFIER_LOCAL) as f:
            sandbox.files.write(VERIFIER_REMOTE, f.read())

        # ── Tests with LibreOffice NOT running ──
        test_help(sandbox)
        test_errors_lo_not_running(sandbox)
        test_errors_bad_args(sandbox)

        # ── Create ODT test file ──
        print("\n  Creating ODT test file...")
        sandbox.files.write("/home/user/create_odt.py", CREATE_ODT_SCRIPT)
        r = run_shell(sandbox, "python3 /home/user/create_odt.py")
        odt_ok = "OK" in r.stdout
        check("ODT file created", odt_ok, f"stdout={r.stdout[:100]} stderr={r.stderr[:100]}")
        if odt_ok:
            test_odt_file_parsing(sandbox)

        # ── Create DOCX test file ──
        print("\n  Creating DOCX test file...")
        sandbox.files.write("/home/user/create_docx.py", CREATE_DOCX_SCRIPT)
        r = run_shell(sandbox, "python3 /home/user/create_docx.py", timeout=120)
        docx_ok = "OK" in r.stdout
        check("DOCX file created", docx_ok, f"stdout={r.stdout[:100]} stderr={r.stderr[:100]}")

        # ── Launch LibreOffice with DOCX file ──
        print("\nLaunching LibreOffice Writer with DOCX file...")
        lo_ready = launch_libreoffice(sandbox, TEST_DOCX_PATH)
        if not lo_ready:
            print("  WARNING: LibreOffice UNO not ready -- UNO tests may fail")

        # ── UNO tests ──
        test_uno_basic(sandbox)
        test_uno_paragraphs(sandbox)
        test_uno_tables(sandbox)
        test_uno_misc(sandbox)
        test_checks_positive(sandbox)
        test_checks_negative(sandbox)
        test_all_commands_return_json(sandbox)

    except Exception:
        traceback.print_exc()
        failed += 1
        errors.append(f"Unhandled exception: {traceback.format_exc()}")

    finally:
        sandbox.kill()
        print("\nSandbox killed.")

    # ── Summary ──
    print("\n" + "=" * 60)
    print(f"Results: {passed} passed, {failed} failed, {passed + failed} total")
    print("=" * 60)

    if errors:
        print("\nFailures:")
        for e in errors:
            print(f"  - {e}")

    sys.exit(1 if failed > 0 else 0)


if __name__ == "__main__":
    main()
